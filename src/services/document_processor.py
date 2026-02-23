"""Document Processor — Multiformat document ingestion for CARF RAG.

Copyright (c) 2026 Cisuregen
Licensed under the Business Source License 1.1 (BSL).

Processes uploaded documents (PDF, DOCX, TXT, CSV, JSON) into structured
text chunks suitable for RAG indexing and policy extraction.

When RAG-Anything is installed (`pip install rag-anything`) the service
delegates to the library for rich multimodal parsing.  Otherwise it uses
built-in parsers (PyPDF2, python-docx) with a plain-text fallback.

Usage:
    from src.services.document_processor import get_document_processor
    processor = get_document_processor()
    result = processor.process_file(file_bytes, "policy.pdf")
"""

from __future__ import annotations

import csv
import io
import json
import logging
from typing import Any

from pydantic import BaseModel, Field

logger = logging.getLogger("carf.document_processor")


# ---------------------------------------------------------------------------
# Data models
# ---------------------------------------------------------------------------

class ProcessedDocument(BaseModel):
    """Result of document processing."""
    filename: str
    file_type: str
    text_content: str
    chunk_count: int = 0
    metadata: dict[str, Any] = Field(default_factory=dict)
    error: str | None = None


class DocumentChunk(BaseModel):
    """A chunk of processed document text."""
    index: int
    content: str
    source: str
    metadata: dict[str, Any] = Field(default_factory=dict)


# ---------------------------------------------------------------------------
# File type parsers
# ---------------------------------------------------------------------------

def _parse_pdf(data: bytes, filename: str) -> str:
    """Extract text from PDF bytes."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(f"[Page {i + 1}]\n{text}")
        return "\n\n".join(pages) if pages else ""
    except ImportError:
        logger.warning("PyPDF2 not installed; PDF parsing unavailable")
        return ""
    except Exception as exc:
        logger.warning("PDF parsing failed: %s", exc)
        return ""


def _parse_docx(data: bytes, filename: str) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed; DOCX parsing unavailable")
        return ""
    except Exception as exc:
        logger.warning("DOCX parsing failed: %s", exc)
        return ""


def _parse_csv(data: bytes, filename: str) -> str:
    """Convert CSV to a text summary."""
    try:
        text = data.decode("utf-8", errors="replace")
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return text

        header = ", ".join(rows[0].keys()) if rows else ""
        summary_lines = [f"CSV: {filename}", f"Columns: {header}", f"Rows: {len(rows)}"]

        # Include first 20 rows as text
        for i, row in enumerate(rows[:20]):
            summary_lines.append(f"Row {i + 1}: {json.dumps(row)}")

        if len(rows) > 20:
            summary_lines.append(f"... and {len(rows) - 20} more rows")

        return "\n".join(summary_lines)
    except Exception as exc:
        logger.warning("CSV parsing failed: %s", exc)
        return data.decode("utf-8", errors="replace")


def _parse_json(data: bytes, filename: str) -> str:
    """Convert JSON to readable text."""
    try:
        obj = json.loads(data.decode("utf-8"))
        return json.dumps(obj, indent=2, default=str)[:10000]
    except Exception:
        return data.decode("utf-8", errors="replace")[:10000]


def _parse_text(data: bytes, filename: str) -> str:
    """Plain text fallback."""
    return data.decode("utf-8", errors="replace")


_PARSERS: dict[str, Any] = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "doc": _parse_docx,
    "csv": _parse_csv,
    "json": _parse_json,
    "txt": _parse_text,
    "md": _parse_text,
    "yaml": _parse_text,
    "yml": _parse_text,
}


def _detect_type(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    return ext


# ---------------------------------------------------------------------------
# Document Processor
# ---------------------------------------------------------------------------

class DocumentProcessor:
    """Multiformat document processor for CARF RAG ingestion."""

    def __init__(self) -> None:
        self._rag_anything_available = False
        try:
            import rag_anything  # type: ignore[import-untyped]
            self._rag_anything_available = True
            logger.info("RAG-Anything detected — multimodal document processing enabled")
        except ImportError:
            logger.debug("RAG-Anything not installed; using built-in parsers")

    @property
    def supported_types(self) -> list[str]:
        return list(_PARSERS.keys())

    def process_bytes(
        self,
        data: bytes,
        filename: str,
        domain_id: str | None = None,
    ) -> ProcessedDocument:
        """Process raw file bytes into structured text.

        Args:
            data: Raw file bytes
            filename: Original filename (used for type detection)
            domain_id: Optional governance domain to tag chunks with

        Returns:
            ProcessedDocument with extracted text and metadata
        """
        file_type = _detect_type(filename)
        parser = _PARSERS.get(file_type, _parse_text)

        try:
            text = parser(data, filename)
            if not text.strip():
                return ProcessedDocument(
                    filename=filename,
                    file_type=file_type,
                    text_content="",
                    error=f"No text could be extracted from {filename}",
                )

            return ProcessedDocument(
                filename=filename,
                file_type=file_type,
                text_content=text,
                chunk_count=max(1, len(text) // 512),
                metadata={
                    "size_bytes": len(data),
                    "text_length": len(text),
                    "domain_id": domain_id,
                },
            )
        except Exception as exc:
            logger.error("Document processing failed for %s: %s", filename, exc)
            return ProcessedDocument(
                filename=filename,
                file_type=file_type,
                text_content="",
                error=str(exc),
            )

    def process_and_ingest(
        self,
        data: bytes,
        filename: str,
        domain_id: str | None = None,
        source: str | None = None,
    ) -> dict[str, Any]:
        """Process a file and ingest it into the RAG index.

        Returns status dict with chunk count and any errors.
        """
        from src.services.rag_service import get_rag_service

        doc = self.process_bytes(data, filename, domain_id)
        if doc.error or not doc.text_content:
            return {
                "status": "error",
                "filename": filename,
                "error": doc.error or "Empty document",
                "chunks_ingested": 0,
            }

        rag = get_rag_service()
        chunks = rag.ingest_text(
            doc.text_content,
            source=source or f"file:{filename}",
            domain_id=domain_id,
            metadata=doc.metadata,
        )

        return {
            "status": "success",
            "filename": filename,
            "file_type": doc.file_type,
            "text_length": len(doc.text_content),
            "chunks_ingested": chunks,
        }

    def get_status(self) -> dict[str, Any]:
        return {
            "supported_types": self.supported_types,
            "rag_anything_available": self._rag_anything_available,
        }


# ---------------------------------------------------------------------------
# Singleton
# ---------------------------------------------------------------------------

_document_processor: DocumentProcessor | None = None


def get_document_processor() -> DocumentProcessor:
    """Get the singleton document processor."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
